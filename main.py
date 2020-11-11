import docx
import re
import csv
regex = re.compile(r"\[.* ]gm")
document = docx.Document('text.docx')

words_list =[]
counter = 0

for paragraph in document.paragraphs:
    fin = ''
    definition = ''
    example = ''
    use = ''
    
    try:
        translation = re.findall('\[.*]', paragraph.text)[0][1:-1]
    except:
        print(paragraph.text)
    for run in paragraph.runs:
        if run.bold:
            fin = fin + run.text
        if not run.italic and not run.bold:
            definition = definition + run.text
        if run.italic:
            example = example + run.text

    definition = re.sub('\[.*]', '', definition)
    words_list.append([fin, definition, example, translation])



with open("dict.csv","w+") as my_csv:
    csvWriter = csv.writer(my_csv,delimiter=',')
    csvWriter.writerows(words_list)











# for paragraph in document.paragraphs:
#     fin = ''
#     definition =''
#     example = ''
#
#     for run in paragraph.runs:
#         if run.bold:
#             line = line + run.text
#
#     print(line)
#
#
#
#
# for paragraph in document.paragraphs:
#     line = ''
#     for run in paragraph.runs:
#
#         if run.italic:
#             line = line + run.text
#     print(line)
#
# for paragraph in document.paragraphs:
#     translation = re.findall('\[.*]', paragraph.text)[0][1:-1]
#     print(translation)
#
# for paragraph in document.paragraphs:
#     line = ''
#     for run in paragraph.runs:
#         if not run.italic and not run.bold:
#             line = line + run.text
#     line = re.sub('\[.*]', '', line)
#     print(line)




